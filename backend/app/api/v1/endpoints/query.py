# backend/app/api/v1/endpoints/query.py
from fastapi import APIRouter, Depends, HTTPException
from fastapi.responses import FileResponse
from sqlalchemy.orm import Session
from sqlalchemy import text, inspect
from app.core.database import get_db
from app.core.deps import tutti_i_ruoli
from app.models.persona import Persona
import requests
import json
from datetime import datetime, date
from decimal import Decimal
import uuid
import os
import tempfile
import re

router = APIRouter()

# Whitelist tabelle per ruolo
ALLOWED_TABLES_BY_ROLE = {
    'superadmin': ['*'],
    'amministrativo': ['spesa', 'rimborso_spesa', 'budget', 'sal', 'progetto', 'allocazione', 'persona', 'timesheet'],
    'ricercatore': ['timesheet', 'spesa', 'rimborso_spesa', 'allocazione', 'progetto'],
    'direttore_generale': ['*'],
    'responsabile_scientifico': ['rimborso_spesa', 'sal', 'spesa', 'timesheet', 'progetto'],
    'monitor': ['timesheet', 'spesa', 'rimborso_spesa', 'sal', 'budget'],
    'direttore_istituto': ['timesheet', 'allocazione', 'persona'],
}

def get_db_schema_dynamically(db: Session) -> str:
    """Legge lo schema reale dal database"""
    try:
        inspector = inspect(db.get_bind())
        tables = inspector.get_table_names()

        schema_str = "Tabelle disponibili nel database:\n"
        for table in tables:
            columns = inspector.get_columns(table)
            col_names = [f"{col['name']}" for col in columns]
            schema_str += f"- {table}: {', '.join(col_names[:8])}"
            if len(col_names) > 8:
                schema_str += f" (+{len(col_names)-8} more)"
            schema_str += "\n"

        schema_str += "\nNOTE IMPORTANTI:\n"
        schema_str += "- progetto.codice: codice tecnico (es: HADEA_001)\n"
        schema_str += "- progetto.acronimo: acronimo progetto (es: WASTE) - per ricerche per nome\n"

        return schema_str
    except Exception as e:
        # Fallback a schema statico se lettura dinamica fallisce
        return """Tabelle disponibili:
- persona: id, nome, cognome, username, email, ruolo, data_inizio_servizio, attivo, dipartimento_id
- progetto: id, codice, acronimo, stato, amministrativo_id, direttore_generale_id
- allocazione: id, persona_id, progetto_id, ruolo, is_pi, data_inizio, data_fine
- timesheet: id, persona_id, progetto_id, settimana, ore, stato
- spesa: id, persona_id, progetto_id, importo, categoria, data, stato
- rimborso_spesa: id, persona_id, progetto_id, importo, data_richiesta, stato
- sal: id, progetto_id, numero, stato, data_invio
- budget: id, progetto_id, voce_costo, importo_allocato
"""

def pulisci_sql(sql: str) -> str:
    """Pulisce SQL generato da Ollama"""
    # Rimuovi backticks (Ollama usa backticks, PostgreSQL usa doppi apici)
    sql = sql.replace('`', '')

    # Rimuovi punto e virgola nel mezzo della query (ma non alla fine)
    sql = re.sub(r';\s+(?!$)', ' ', sql)

    # Rimuovi markdown code blocks se presenti
    if sql.startswith('```'):
        sql = sql.split('```')[1].replace('sql', '').strip()

    # Normalizza whitespace
    sql = ' '.join(sql.split())

    # Assicurati che finisca con LIMIT se non ha ORDER BY
    if 'LIMIT' not in sql.upper():
        sql = sql.rstrip(';') + ' LIMIT 1000'

    return sql.strip()

def valida_sql_syntax(sql: str) -> tuple[bool, str]:
    """Valida sintassi SQL base"""
    sql_upper = sql.upper().strip()

    # Rimuovi commenti SQL (-- e /* */)
    sql_no_comments = re.sub(r'--.*?$', '', sql_upper, flags=re.MULTILINE)
    sql_no_comments = re.sub(r'/\*.*?\*/', '', sql_no_comments, flags=re.DOTALL)
    sql_no_comments = sql_no_comments.strip()

    # Vieta operazioni pericolose
    dangerous = ['DROP', 'DELETE', 'UPDATE', 'INSERT', 'ALTER', 'TRUNCATE', 'CREATE']
    for kw in dangerous:
        if kw in sql_no_comments:
            return False, f"Operazione '{kw}' non permessa"

    # Verifica che sia un SELECT
    if not sql_no_comments.startswith('SELECT'):
        return False, f"Solo query SELECT sono permesse (trovato: {sql_no_comments[:50]})"

    # Verifica parentesi bilanciate
    if sql.count('(') != sql.count(')'):
        return False, "Parentesi non bilanciate"

    return True, "OK"

def valida_sql_con_explain(db: Session, sql: str) -> tuple[bool, str]:
    """Valida SQL eseguendo EXPLAIN (PostgreSQL)"""
    try:
        # Rimuovi punto e virgola finale
        sql_clean = sql.rstrip(';')
        explain_sql = f"EXPLAIN {sql_clean}"
        db.execute(text(explain_sql))
        return True, "SQL valido"
    except Exception as e:
        return False, str(e)

def estrai_tabelle_sql(sql: str) -> list:
    """Estrae nomi tabelle da SQL"""
    pattern = r'(?:FROM|JOIN)\s+(\w+)'
    matches = re.findall(pattern, sql, re.IGNORECASE)
    return [m.lower() for m in matches]

def valida_tabelle_per_ruolo(sql: str, ruolo: str) -> bool:
    """Valida che tabelle siano allowed per il ruolo"""
    tabelle_sql = estrai_tabelle_sql(sql)
    allowed = ALLOWED_TABLES_BY_ROLE.get(ruolo, [])

    if '*' in allowed:
        return True

    for tabella in tabelle_sql:
        if tabella not in allowed and tabella not in ['pg_stat_statements']:  # Ignora tabelle di sistema
            return False

    return True

def aggiungi_filtri_sicurezza(sql: str, ruolo: str, utente: Persona) -> str:
    """Aggiunge filtri di sicurezza obbligatori"""

    if ruolo == 'ricercatore':
        # Ricercatore vede solo i SUOI dati
        if 'WHERE' in sql.upper():
            return sql + f" AND persona.id = '{utente.id}'"
        elif 'FROM persona' in sql.lower():
            return sql + f" WHERE persona.id = '{utente.id}'"

    elif ruolo == 'amministrativo':
        # Admin vede solo il suo progetto
        if 'WHERE' in sql.upper():
            return sql + f" AND progetto.amministrativo_id = '{utente.id}'"
        elif 'FROM progetto' in sql.lower():
            return sql + f" WHERE progetto.amministrativo_id = '{utente.id}'"

    elif ruolo == 'responsabile_scientifico':
        # RS vede solo i suoi progetti
        if 'WHERE' in sql.upper():
            return sql + f" AND progetto.id IN (SELECT progetto_id FROM allocazione WHERE persona_id = '{utente.id}')"
        elif 'FROM progetto' in sql.lower():
            return sql + f" WHERE progetto.id IN (SELECT progetto_id FROM allocazione WHERE persona_id = '{utente.id}')"

    return sql

def converti_risultati_json(column_names: list, data: list) -> list:
    """Converte risultati in formato JSON-serializzabile"""
    # Identifica colonne che sono importi
    importo_keywords = ['importo', 'costo', 'prezzo', 'totale', 'spesa', 'budget', 'allocato', 'disponibile']
    importo_columns = {i for i, col in enumerate(column_names)
                      if any(kw in col.lower() for kw in importo_keywords)}

    risultati = []
    for row in data:
        row_dict = {}
        for col_idx, (col_name, value) in enumerate(zip(column_names, row)):
            if value is None:
                row_dict[col_name] = None
            elif isinstance(value, (uuid.UUID, date, datetime)):
                row_dict[col_name] = str(value)
            elif isinstance(value, Decimal) or isinstance(value, (int, float)):
                # Se è un importo, formatta come euro italiano
                if col_idx in importo_columns:
                    num_value = float(value)
                    row_dict[col_name] = f"€ {num_value:,.2f}".replace(',', 'X').replace('.', ',').replace('X', '.')
                else:
                    row_dict[col_name] = float(value) if isinstance(value, Decimal) else value
            elif isinstance(value, bool):
                row_dict[col_name] = value
            else:
                row_dict[col_name] = str(value) if not isinstance(value, (int, float, str)) else value
        risultati.append(row_dict)
    return risultati

def genera_sql_con_ollama(domanda: str, schema: str) -> str:
    """Genera SQL usando Ollama locale"""

    prompt = f"""
    Sei un esperto di SQL PostgreSQL.

    {schema}

    Converti questa domanda naturale in SQL query PostgreSQL valida:
    "{domanda}"

    REGOLE CRITICHE:
    1. Ritorna SOLO la query SQL (niente markdown, niente backticks)
    2. USA SOLO tabelle dalla lista sopra
    3. Usa SELECT ... FROM ... WHERE ... JOIN ...
    4. NON aggiungere backticks (`) - PostgreSQL usa doppi apici (")
    5. NON usare punto e virgola (;) nel mezzo
    6. AGGIUNGI sempre LIMIT 1000 alla fine
    7. Niente operazioni pericolose (DROP, DELETE, UPDATE, INSERT, ALTER)

    Esempio:
    Domanda: "Mostra i primi 10 progetti"
    SQL: SELECT id, codice, acronimo, stato FROM progetto LIMIT 10

    SQL Query:
    """

    try:
        response = requests.post(
            'http://host.docker.internal:11434/api/generate',
            json={
                "model": "mistral",
                "prompt": prompt,
                "stream": False,
                "temperature": 0.2  # Bassa temp per stabilità
            },
            timeout=60
        )

        sql_query = response.json()['response'].strip()
        return sql_query

    except Exception as e:
        raise Exception(f"Ollama error: {str(e)}")

@router.post("/naturale")
def query_naturale(
    payload: dict,
    db: Session = Depends(get_db),
    utente: Persona = Depends(tutti_i_ruoli),
):
    """
    Interrogazione DB in linguaggio naturale con Ollama

    ⚠️ SOLO SUPERADMIN - Funzionalità in Beta

    Body: {
        "domanda": "Quante ore ha lavorato Giuseppe?",
        "formato": "json"
    }
    """

    # Controllo accesso: solo superadmin
    if utente.ruolo != 'superadmin':
        raise HTTPException(
            status_code=403,
            detail="Query IA disponibile solo per superadmin (funzionalità Beta)"
        )

    domanda = payload.get("domanda")
    formato = payload.get("formato", "json")

    if not domanda or not domanda.strip():
        raise HTTPException(status_code=400, detail="Domanda mancante")

    try:
        # 1. Leggi schema reale dal DB
        schema = get_db_schema_dynamically(db)

        # 2. Genera SQL con Ollama
        sql_query = genera_sql_con_ollama(domanda, schema)

        # 3. Pulisci SQL
        sql_pulito = pulisci_sql(sql_query)

        # 4. Valida sintassi SQL
        is_valid, msg = valida_sql_syntax(sql_pulito)
        if not is_valid:
            raise HTTPException(status_code=400, detail=f"SQL invalido: {msg}")

        # 5. Valida tabelle per ruolo
        if not valida_tabelle_per_ruolo(sql_pulito, utente.ruolo):
            raise HTTPException(status_code=403, detail="Non hai accesso a queste tabelle")

        # 6. Aggiungi filtri di sicurezza
        sql_sicuro = aggiungi_filtri_sicurezza(sql_pulito, utente.ruolo, utente)

        # 7. Valida con EXPLAIN PLAN
        is_valid, msg = valida_sql_con_explain(db, sql_sicuro)
        if not is_valid:
            raise HTTPException(status_code=400, detail=f"SQL non eseguibile: {msg}")

        # 8. Esegui query
        result = db.execute(text(sql_sicuro))
        rows = result.fetchall()

        # 9. Formatta dati
        column_names = list(result.keys())
        data = converti_risultati_json(column_names, rows)

        # 10. Ritorna risultati
        if formato == "json":
            return {
                "domanda": domanda,
                "sql_generato": sql_sicuro,
                "conteggio": len(data),
                "risultati": data,
                "note": "Query eseguita con successo"
            }

        elif formato == "excel":
            file_path = genera_excel(data, domanda, utente.username)
            return FileResponse(
                file_path,
                media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                filename=f"query_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"
            )

        elif formato == "pdf":
            file_path = genera_pdf(data, domanda, utente.username)
            return FileResponse(
                file_path,
                media_type="application/pdf",
                filename=f"query_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
            )

        elif formato == "word":
            file_path = genera_word(data, domanda, utente.username)
            return FileResponse(
                file_path,
                media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                filename=f"query_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
            )

    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Errore: {str(e)}")

def genera_excel(data: list, domanda: str, username: str) -> str:
    """Genera file Excel"""
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment

    file_id = str(uuid.uuid4())
    file_path = os.path.join(tempfile.gettempdir(), f"query_{file_id}.xlsx")

    wb = Workbook()
    ws = wb.active
    ws.title = "Risultati"

    # Intestazione
    ws['A1'] = "INTERROGAZIONE DATABASE"
    ws['A1'].font = Font(size=14, bold=True, color="FFFFFF")
    ws['A1'].fill = PatternFill(start_color="1E5AA0", end_color="1E5AA0", fill_type="solid")
    ws.merge_cells('A1:D1')

    # Metadata
    ws['A2'] = f"Domanda: {domanda}"
    ws['A3'] = f"Utente: {username}"
    ws['A4'] = f"Data: {datetime.now().strftime('%d/%m/%Y %H:%M')}"
    ws['A5'] = f"Risultati: {len(data)}"

    # Dati
    if data:
        headers = list(data[0].keys())
        for col, header in enumerate(headers, 1):
            cell = ws.cell(row=7, column=col)
            cell.value = header
            cell.font = Font(bold=True, color="FFFFFF")
            cell.fill = PatternFill(start_color="1E5AA0", end_color="1E5AA0", fill_type="solid")

        for row_idx, row_data in enumerate(data, 8):
            for col_idx, value in enumerate(headers, 1):
                ws.cell(row=row_idx, column=col_idx).value = row_data.get(value)

    for col in ws.columns:
        max_length = max(len(str(cell.value or "")) for cell in col)
        ws.column_dimensions[col[0].column_letter].width = min(max_length + 2, 50)

    wb.save(file_path)
    return file_path

def genera_pdf(data: list, domanda: str, username: str) -> str:
    """Genera file PDF"""
    from reportlab.lib.pagesizes import A4
    from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.lib import colors

    file_id = str(uuid.uuid4())
    file_path = os.path.join(tempfile.gettempdir(), f"query_{file_id}.pdf")

    doc = SimpleDocTemplate(file_path, pagesize=A4)
    story = []
    styles = getSampleStyleSheet()

    story.append(Paragraph(f"Interrogazione: {domanda}", styles['Heading1']))
    story.append(Spacer(1, 0.2*inch))

    if data:
        headers = list(data[0].keys())
        table_data = [headers] + [[str(row.get(h, "")) for h in headers] for row in data[:100]]

        table = Table(table_data, colWidths=[6.5*inch / len(headers)] * len(headers))
        table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor("#1E5AA0")),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTSIZE', (0, 0), (-1, -1), 8),
            ('GRID', (0, 0), (-1, -1), 1, colors.grey),
        ]))
        story.append(table)

    doc.build(story)
    return file_path

def genera_word(data: list, domanda: str, username: str) -> str:
    """Genera file Word"""
    from docx import Document
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

    file_id = str(uuid.uuid4())
    file_path = os.path.join(tempfile.gettempdir(), f"query_{file_id}.docx")

    doc = Document()
    title = doc.add_heading(f"Interrogazione: {domanda}", 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.add_paragraph(f"Utente: {username}")
    doc.add_paragraph(f"Data: {datetime.now().strftime('%d/%m/%Y %H:%M')}")
    doc.add_paragraph(f"Risultati: {len(data)}")

    if data:
        headers = list(data[0].keys())
        table = doc.add_table(rows=1, cols=len(headers))

        for i, header in enumerate(headers):
            table.rows[0].cells[i].text = header

        for row_data in data[:100]:
            row_cells = table.add_row().cells
            for i, header in enumerate(headers):
                row_cells[i].text = str(row_data.get(header, ""))

    doc.save(file_path)
    return file_path
